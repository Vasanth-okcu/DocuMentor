import os
import pytesseract
from PIL import Image
import docx
import PyPDF2
import fitz  # PyMuPDF
import pdfplumber
from io import BytesIO

def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()
    
def extract_text_from_pdf(file_path):
    text = ""
    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text("text").strip() + " "  # Strip extra whitespace and line breaks
    return text.strip()

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        # Add space instead of newline for each paragraph to avoid unwanted line breaks
        text += para.text.strip() + " "  # Strip extra spaces from paragraphs
    return text.strip()

def extract_images_from_pdf(file_path):
    images = []
    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        img_list = page.get_images(full=True)
        for img_index, img in enumerate(img_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_data = base_image["image"]
            image = Image.open(BytesIO(image_data))
            images.append(image)
    return images

def extract_images_from_docx(file_path):
    images = []
    doc = docx.Document(file_path)
    rels = doc.part.rels
    for rel in rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image = Image.open(BytesIO(image_data))
            images.append(image)
    return images

def extract_text_from_image(image):
    return pytesseract.image_to_string(image).strip()

def extract_page_data(file_path):
    file_extension = os.path.splitext(file_path)[-1].lower()
    page_data = {}

    if file_extension == '.txt':
        text_data = extract_text_from_txt(file_path)
        page_data['page_1'] = {'text': text_data, 'images': []}

    elif file_extension == '.pdf':
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            text_data = doc.load_page(page_num).get_text("text").strip()
            images = []
            img_list = doc.load_page(page_num).get_images(full=True)
            for img in img_list:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_data = base_image["image"]
                image = Image.open(BytesIO(image_data))
                images.append(image)
            page_data[f'page_{page_num + 1}'] = {'text': text_data, 'images': images}

    elif file_extension == '.docx':
        doc = docx.Document(file_path)
        text_data = ""
        for para in doc.paragraphs:
            text_data += para.text.strip() + " "
        images = extract_images_from_docx(file_path)
        page_data['page_1'] = {'text': text_data.strip(), 'images': images}

    else:
        raise ValueError("Unsupported file format")
    
    return page_data

def extract_data(file_path):
    data = extract_page_data(file_path)
    result = []

    for page, content in data.items():
        page_info = {}
        page_info['page'] = page
        page_info['text'] = content['text']
        page_info['images'] = []
        
        # Extract text from images if available
        for i, img in enumerate(content['images']):
            text_on_image = extract_text_from_image(img)
            page_info['images'].append({
                'image': img,  # The actual image object
                'text_from_image': text_on_image  # OCR text from the image
            })
        
        result.append(page_info)
    
    return result


file_path = r"C:\\Users\\Vasanth\\Desktop\\which chart when.pdf"
output = extract_data(file_path)

for page_info in output:
    print(f"Page: {page_info['page']}")
    print(f"Text:\n{page_info['text']}")
    print(f"Number of Images: {len(page_info['images'])}")
    for img_info in page_info['images']:
        print(f"Image Text: {img_info['text_from_image']}")

